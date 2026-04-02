"""engine.py — The RenderEngine: generate documents and code from RDF.

Architecture
------------
1. Load RDF data from multiple sources into a single graph.
2. Load (or define inline) a RenderSpec — an RDF description of what
   to generate, which template to use, and which SPARQL queries supply
   the template variables.
3. The engine reads the RenderSpec from the graph, executes each
   DataBinding's SPARQL query, feeds results to the Jinja2 template,
   and dispatches the output to the appropriate renderer.

The entire generation pipeline is described in RDF.  The RenderSpec,
Template, and DataBindings are all queryable graph entities.

Usage
-----
::

    from rdf_gen import RenderEngine, load_ttl

    engine = RenderEngine()
    engine.load_data(my_data_ttl)       # domain data in TTL
    engine.load_spec(my_spec_ttl)       # render spec in TTL
    engine.render("urn:spec:my-doc")    # execute the spec
    engine.render_all()                 # execute all specs in the graph
"""

from __future__ import annotations

import os
from io import StringIO
from pathlib import Path
from typing import Any

from rdflib import RDFS, XSD, Graph, URIRef

from .namespaces import GEN, SPARQL_PREFIXES, TTL_PREFIXES
from .queries import execute_construct, execute_select


class RenderEngine:
    """Generate documents and code from RDF data using graph-described specs.

    The engine maintains a single merged RDF graph containing both domain
    data and generation specifications (RenderSpecs, Templates, DataBindings).
    """

    def _bind_prefixes(self):
        from rdflib.namespace import SH

        self.graph.bind("gen", GEN)
        self.graph.bind("rdfs", RDFS)
        self.graph.bind("sh", SH)
        self.graph.bind("xsd", XSD)

    def __init__(self, template_dirs: list[str] | None = None):
        """Parameters
        ----------
        template_dirs : list[str], optional
            Directories to search for template files referenced by
            ``gen:templatePath``.  The current directory is always included.
        """
        self.graph = Graph()
        self._bind_prefixes()
        self._template_dirs = ["."]
        if template_dirs:
            self._template_dirs.extend(template_dirs)
        self._jinja_env = None

    # ------------------------------------------------------------------
    # Data loading
    # ------------------------------------------------------------------

    def load_data(self, ttl: str) -> None:
        """Parse TTL string into the data graph."""
        full = TTL_PREFIXES + "\n" + ttl
        self.graph.parse(StringIO(full), format="turtle")

    def load_file(self, path: str, fmt: str = "turtle") -> None:
        """Parse a file into the data graph."""
        self.graph.parse(path, format=fmt)

    def load_spec(self, ttl: str) -> None:
        """Parse a RenderSpec (with templates and bindings) into the graph.
        This is the same graph as the data — specs and data coexist.
        """
        self.load_data(ttl)

    def load_graph(self, g: Graph) -> None:
        """Merge an existing rdflib Graph into the engine's graph."""
        for triple in g:
            self.graph.add(triple)

    # ------------------------------------------------------------------
    # Spec discovery
    # ------------------------------------------------------------------

    def list_specs(self) -> list[dict]:
        """Return all RenderSpecs in the graph."""
        query = (
            SPARQL_PREFIXES
            + """
            SELECT ?spec ?title ?description ?format ?outPath
            WHERE {
                ?spec a gen:RenderSpec .
                OPTIONAL { ?spec gen:title ?title }
                OPTIONAL { ?spec gen:description ?description }
                OPTIONAL {
                    ?spec gen:usesTemplate ?tmpl .
                    ?tmpl gen:outputFormat ?format
                }
                OPTIONAL { ?spec gen:outputPath ?outPath }
            }
            ORDER BY ?spec
        """
        )
        return execute_select(self.graph, query, result_type="list")

    def _scalar(self, subject: str, predicate) -> str | None:
        for _, _, o in self.graph.triples((URIRef(subject), predicate, None)):
            return str(o)
        return None

    def _read_spec(self, spec_iri: str) -> dict:
        """Read a complete RenderSpec from the graph."""
        spec_uri = URIRef(spec_iri)

        # Template
        tmpl_query = (
            SPARQL_PREFIXES
            + f"""
            SELECT ?tmpl ?body ?path ?format ?ext
            WHERE {{
                <{spec_iri}> gen:usesTemplate ?tmpl .
                OPTIONAL {{ ?tmpl gen:templateBody ?body }}
                OPTIONAL {{ ?tmpl gen:templatePath ?path }}
                OPTIONAL {{ ?tmpl gen:outputFormat ?format }}
                OPTIONAL {{ ?tmpl gen:outputExtension ?ext }}
            }}
        """
        )
        tmpl_rows = execute_select(self.graph, tmpl_query)
        if not tmpl_rows:
            raise ValueError(f"RenderSpec {spec_iri} has no template.")

        tmpl = tmpl_rows[0]

        # Bindings
        bind_query = (
            SPARQL_PREFIXES
            + f"""
            SELECT ?binding ?varName ?query ?constructQuery ?resultType ?default
            WHERE {{
                <{spec_iri}> gen:hasBinding ?binding .
                ?binding gen:variableName ?varName .
                OPTIONAL {{ ?binding gen:sparqlQuery ?query }}
                OPTIONAL {{ ?binding gen:sparqlConstructQuery ?constructQuery }}
                OPTIONAL {{ ?binding gen:resultType ?resultType }}
                OPTIONAL {{ ?binding gen:defaultValue ?default }}
            }}
            ORDER BY ?varName
        """
        )
        bindings = execute_select(self.graph, bind_query)

        # Sections (ordered)
        section_query = (
            SPARQL_PREFIXES
            + f"""
            SELECT ?section ?order ?sTitle ?varName ?query ?resultType
            WHERE {{
                <{spec_iri}> gen:hasSection ?section .
                ?section gen:sectionOrder ?order .
                OPTIONAL {{ ?section gen:sectionTitle ?sTitle }}
                OPTIONAL {{
                    ?section gen:sectionBinding ?binding .
                    ?binding gen:variableName ?varName .
                    ?binding gen:sparqlQuery ?query .
                    OPTIONAL {{ ?binding gen:resultType ?resultType }}
                }}
            }}
            ORDER BY ?order
        """
        )
        sections = execute_select(self.graph, section_query)

        # Metadata
        title = self._scalar(spec_iri, GEN.title) or spec_iri
        description = self._scalar(spec_iri, GEN.description) or ""
        output_path = self._scalar(spec_iri, GEN.outputPath)

        return {
            "iri": spec_iri,
            "title": title,
            "description": description,
            "output_path": output_path,
            "template": tmpl,
            "bindings": bindings,
            "sections": sections,
        }

    # ------------------------------------------------------------------
    # Template resolution
    # ------------------------------------------------------------------

    def _get_jinja_env(self):
        """Lazy-init the Jinja2 environment."""
        if self._jinja_env is None:
            from jinja2 import BaseLoader, Environment, FileSystemLoader

            if any(os.path.isdir(d) for d in self._template_dirs):
                loader = FileSystemLoader(self._template_dirs)
            else:
                loader = BaseLoader()
            self._jinja_env = Environment(
                loader=loader,
                keep_trailing_newline=True,
                trim_blocks=True,
                lstrip_blocks=True,
            )
            # Custom filters
            self._jinja_env.filters["shorten"] = lambda s: (
                s.rsplit("/", 1)[-1].rsplit(":", 1)[-1].rsplit("#", 1)[-1]
                if isinstance(s, str)
                else str(s)
            )
            self._jinja_env.filters["uri_local"] = self._jinja_env.filters["shorten"]
        return self._jinja_env

    def _resolve_template(self, tmpl_info: dict) -> jinja2.Template:
        """Resolve a template from inline body or file path."""
        env = self._get_jinja_env()

        body = tmpl_info.get("body")
        path = tmpl_info.get("path")

        if body:
            return env.from_string(body)
        if path:
            return env.get_template(path)

        raise ValueError("Template has neither templateBody nor templatePath.")

    # ------------------------------------------------------------------
    # Binding execution
    # ------------------------------------------------------------------

    def _execute_bindings(self, bindings: list[dict]) -> dict[str, Any]:
        """Execute all SPARQL bindings and return a template context dict."""
        context = {}

        for b in bindings:
            var_name = b["varName"]
            result_type = b.get("resultType") or "list"
            default = b.get("default")

            select_q = b.get("query")
            construct_q = b.get("constructQuery")

            if select_q:
                result = execute_select(
                    self.graph,
                    select_q,
                    result_type=result_type,
                    shorten_uris=True,
                )
            elif construct_q:
                result = execute_construct(self.graph, construct_q)
            else:
                result = default

            # Apply default if empty
            if result is None or (isinstance(result, (list, dict)) and len(result) == 0):
                result = default

            context[var_name] = result

        return context

    # ------------------------------------------------------------------
    # Binary format helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _render_docx(markdown_content: str, output_path: str):
        """Convert markdown to docx using python-docx (best-effort)."""
        try:
            from docx import Document

            doc = Document()
            for line in markdown_content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped.startswith("| "):
                    doc.add_paragraph(stripped)  # simplified table
                elif stripped:
                    doc.add_paragraph(stripped)
            docx_path = (
                output_path.replace(".md", ".docx") if output_path.endswith(".md") else output_path
            )
            doc.save(docx_path)
        except ImportError:
            pass  # python-docx not installed

    @staticmethod
    def _render_pdf(markdown_content: str, output_path: str):
        """Placeholder for PDF rendering."""
        pass  # Would use weasyprint, reportlab, or pandoc

    # ------------------------------------------------------------------
    # Rendering
    # ------------------------------------------------------------------

    def render(
        self,
        spec_iri: str,
        output_path: str | None = None,
        extra_context: dict | None = None,
    ) -> str:
        """Execute a single RenderSpec and return the generated content.

        Parameters
        ----------
        spec_iri : str
            IRI of the RenderSpec in the graph.
        output_path : str, optional
            Override the output file path.  If None, uses gen:outputPath
            from the spec, or returns the content without writing.
        extra_context : dict, optional
            Additional variables to pass to the template.

        Returns:
        -------
        str
            The rendered content.
        """
        spec = self._read_spec(spec_iri)
        template = self._resolve_template(spec["template"])

        # Build template context from bindings
        context = self._execute_bindings(spec["bindings"])

        # Add metadata
        context["_title"] = spec["title"]
        context["_description"] = spec["description"]
        context["_spec_iri"] = spec["iri"]
        context["_format"] = spec["template"].get("format", "markdown")

        # Process sections if present
        if spec["sections"]:
            sections = []
            for s in spec["sections"]:
                sec = {"title": s.get("sTitle", ""), "order": s.get("order", 0)}
                if s.get("query"):
                    rt = s.get("resultType") or "list"
                    sec["data"] = execute_select(
                        self.graph,
                        s["query"],
                        result_type=rt,
                        shorten_uris=True,
                    )
                else:
                    sec["data"] = []
                sections.append(sec)
            context["_sections"] = sections

        # Merge extra context
        if extra_context:
            context.update(extra_context)

        # Render
        content = template.render(**context)

        # Write output
        out = output_path or spec.get("output_path")
        if out:
            Path(out).parent.mkdir(parents=True, exist_ok=True)
            Path(out).write_text(content, encoding="utf-8")

        # Post-process for binary formats
        fmt = spec["template"].get("format", "markdown")
        if fmt == "docx" and out:
            self._render_docx(content, out)
        elif fmt == "pdf" and out:
            self._render_pdf(content, out)

        return content

    def render_all(
        self,
        output_dir: str | None = None,
    ) -> dict[str, str]:
        """Execute all RenderSpecs in the graph.

        Returns a dict mapping spec IRI → rendered content.
        """
        specs = self.list_specs()
        results = {}
        for spec in specs:
            iri = spec["spec"]
            out = spec.get("outPath")
            if output_dir and out:
                out = os.path.join(output_dir, os.path.basename(out))
            results[iri] = self.render(iri, output_path=out)
        return results

    # ------------------------------------------------------------------
    # Quick API — render without a formal RenderSpec in the graph
    # ------------------------------------------------------------------

    def render_template(
        self,
        template_str: str,
        queries: dict[str, str],
        result_types: dict[str, str] | None = None,
        extra_context: dict | None = None,
    ) -> str:
        """Quick render: provide a template string and a dict of
        {variable_name: sparql_query} directly, without defining
        a RenderSpec in the graph.

        Parameters
        ----------
        template_str : str
            Jinja2 template content.
        queries : dict[str, str]
            Mapping from template variable names to SPARQL SELECT queries.
        result_types : dict[str, str], optional
            Mapping from variable names to result types (scalar/list/grouped).
        extra_context : dict, optional
            Additional template variables.

        Returns:
        -------
        str
            Rendered content.
        """
        env = self._get_jinja_env()
        template = env.from_string(template_str)

        context = {}
        rt = result_types or {}

        for var_name, query in queries.items():
            rtype = rt.get(var_name, "list")
            context[var_name] = execute_select(
                self.graph,
                query,
                result_type=rtype,
                shorten_uris=True,
            )

        if extra_context:
            context.update(extra_context)

        return template.render(**context)


# ------------------------------------------------------------------
# Convenience function
# ------------------------------------------------------------------


def load_ttl(ttl: str) -> Graph:
    """Parse a TTL string (with auto-prefixed header) into a Graph."""
    full = TTL_PREFIXES + "\n" + ttl
    g = Graph()
    g.parse(StringIO(full), format="turtle")
    return g
